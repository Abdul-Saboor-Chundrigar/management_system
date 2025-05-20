from django.views.generic import FormView
from django.http import HttpResponse
import csv
from io import BytesIO
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from docx import Document
from asset_management.models import Asset
from warehouse_management.models import Warehouse
from vendorescalation_management.models import VendorEscalationManagement
from .forms import ReportForm

class ReportView(FormView):
    template_name = 'report_generation/report_form.html'
    form_class = ReportForm
    success_url = 'report_generation:reports'

    def form_valid(self, form):
        report_type = form.cleaned_data['report_type']
        export_format = form.cleaned_data['export_format']
        action = self.request.POST.get('action', 'view')

        # Get data based on report type with proper related field handling
        if report_type == 'asset':
            queryset = Asset.objects.all().select_related('employee')
            context = self.get_asset_context(queryset)
        elif report_type == 'warehouse':
            queryset = Warehouse.objects.all().select_related('employee')
            context = self.get_warehouse_context(queryset)
        elif report_type == 'vendor':
            queryset = VendorEscalation.objects.all().select_related('employee')
            context = self.get_vendor_context(queryset)

        # Check if we actually got data
        if not context['data']:
            return self.render_to_response(self.get_context_data(
                form=form,
                no_data=True,
                title=context['title'],
                report_type=report_type
            ))

        # Export logic
        if action == 'export':
            if export_format == 'csv':
                return self.export_to_csv(context)
            elif export_format == 'pdf':
                return self.export_to_pdf(context)
            elif export_format == 'xlsx':
                return self.export_to_excel(context)
            elif export_format == 'docx':
                return self.export_to_word(context)

        # HTML view
        return self.render_to_response(self.get_context_data(
            form=form,
            data=context['data'],
            headers=context['headers'],
            report_type=report_type,
            title=context['title'],
            export_format=export_format,
            show_data=True
        ))

    def get_asset_context(self, queryset):
        headers = [
            'Employee', 'Device Type', 'Region', 'Date Assigned',
            'Brand', 'Model', 'Serial #', 'AIMS Tag',
            'LCD Brand', 'LCD Model', 'LCD Serial', 'LCD AIMS Tag',
            'CPU Brand', 'CPU Model', 'CPU Serial', 'CPU AIMS Tag'
        ]
        
        data = []
        for item in queryset:
            employee_name = f"{item.employee.name} ({item.employee.emp_number})" if item.employee else 'N/A'
            region = getattr(item.employee.region, 'name', 'N/A') if item.employee else 'N/A'
            
            data.append({
                'Employee': employee_name,
                'Device Type': str(getattr(item, 'device_type', 'N/A')),
                'Region': region,
                'Date Assigned': item.date_assigned.strftime('%Y-%m-%d') if item.date_assigned else 'N/A',
                'Brand': getattr(item, 'brand', 'N/A'),
                'Model': getattr(item, 'model', 'N/A'),
                'Serial #': getattr(item, 'serial_number', 'N/A'),
                'AIMS Tag': getattr(item, 'aims_tag', 'N/A'),
                'LCD Brand': getattr(item, 'lcd_brand', 'N/A'),
                'LCD Model': getattr(item, 'lcd_model', 'N/A'),
                'LCD Serial': getattr(item, 'lcd_serial', 'N/A'),
                'LCD AIMS Tag': getattr(item, 'lcd_aims_tag', 'N/A'),
                'CPU Brand': getattr(item, 'cpu_brand', 'N/A'),
                'CPU Model': getattr(item, 'cpu_model', 'N/A'),
                'CPU Serial': getattr(item, 'cpu_serial', 'N/A'),
                'CPU AIMS Tag': getattr(item, 'cpu_aims_tag', 'N/A')
            })

        return {
            'data': data,
            'headers': headers,
            'title': 'Employee Asset Records'
        }

    def get_warehouse_context(self, queryset):
        headers = [
            'Date Assigned', 'Employee', 'Device Type', 'Region',
            'Recipient', 'Recipient Email', 'Recipient Cell', 'Dispatcher',
            'IT Email', 'IT Cell', 'Brand', 'Model', 'Serial #', 'AIMS Tag',
            'LCD Brand', 'LCD Model', 'LCD Serial', 'LCD AIMS Tag',
            'CPU Brand', 'CPU Model', 'CPU Serial', 'CPU AIMS Tag'
        ]
        
        data = []
        for item in queryset:
            employee_name = f"{item.employee.name} ({item.employee.emp_number})" if item.employee else 'N/A'
            region = getattr(item.employee.region, 'name', 'N/A') if item.employee else 'N/A'
            
            data.append({
                'Date Assigned': item.date_assigned.strftime('%Y-%m-%d') if item.date_assigned else 'N/A',
                'Employee': employee_name,
                'Device Type': str(getattr(item, 'device_type', 'N/A')),
                'Region': region,
                'Recipient': getattr(item, 'recipient_name', 'N/A'),
                'Recipient Email': getattr(item, 'recipient_email', 'N/A'),
                'Recipient Cell': getattr(item, 'recipient_cell', 'N/A'),
                'Dispatcher': getattr(item, 'dispatcher', 'N/A'),
                'IT Email': getattr(item, 'it_email', 'N/A'),
                'IT Cell': getattr(item, 'it_cell', 'N/A'),
                'Brand': getattr(item, 'brand', 'N/A'),
                'Model': getattr(item, 'model', 'N/A'),
                'Serial #': getattr(item, 'serial_number', 'N/A'),
                'AIMS Tag': getattr(item, 'aims_tag', 'N/A'),
                'LCD Brand': getattr(item, 'lcd_brand', 'N/A'),
                'LCD Model': getattr(item, 'lcd_model', 'N/A'),
                'LCD Serial': getattr(item, 'lcd_serial', 'N/A'),
                'LCD AIMS Tag': getattr(item, 'lcd_aims_tag', 'N/A'),
                'CPU Brand': getattr(item, 'cpu_brand', 'N/A'),
                'CPU Model': getattr(item, 'cpu_model', 'N/A'),
                'CPU Serial': getattr(item, 'cpu_serial', 'N/A'),
                'CPU AIMS Tag': getattr(item, 'cpu_aims_tag', 'N/A')
            })

        return {
            'data': data,
            'headers': headers,
            'title': 'Warehouse Records'
        }

    def get_vendor_context(self, queryset):
        headers = [
            'Date Created', 'Vendor Name', 'Vendor Email', 'Employee',
            'Device Type', 'Brand', 'Model', 'Serial #', 'AIMS Tag',
            'LCD Brand', 'LCD Model', 'LCD Serial', 'LCD AIMS Tag',
            'Issue Details', 'POC Name', 'IT Email', 'POC Cell',
            'Status', 'Closure Date', 'Closure Remarks'
        ]
        
        data = []
        for item in queryset:
            employee_name = f"{item.employee.name} ({item.employee.emp_number})" if item.employee else 'N/A'
            
            data.append({
                'Date Created': item.date_created.strftime('%Y-%m-%d') if item.date_created else 'N/A',
                'Vendor Name': getattr(item, 'vendor_name', 'N/A'),
                'Vendor Email': getattr(item, 'vendor_email', 'N/A'),
                'Employee': employee_name,
                'Device Type': str(getattr(item, 'device_type', 'N/A')),
                'Brand': getattr(item, 'brand', 'N/A'),
                'Model': getattr(item, 'model', 'N/A'),
                'Serial #': getattr(item, 'serial_number', 'N/A'),
                'AIMS Tag': getattr(item, 'aims_tag', 'N/A'),
                'LCD Brand': getattr(item, 'lcd_brand', 'N/A'),
                'LCD Model': getattr(item, 'lcd_model', 'N/A'),
                'LCD Serial': getattr(item, 'lcd_serial', 'N/A'),
                'LCD AIMS Tag': getattr(item, 'lcd_aims_tag', 'N/A'),
                'Issue Details': getattr(item, 'issue_details', 'N/A'),
                'POC Name': getattr(item, 'poc_name', 'N/A'),
                'IT Email': getattr(item, 'it_email', 'N/A'),
                'POC Cell': getattr(item, 'poc_cell', 'N/A'),
                'Status': item.get_status_display() if hasattr(item, 'get_status_display') else 'N/A',
                'Closure Date': item.closure_date.strftime('%Y-%m-%d') if hasattr(item, 'closure_date') and item.closure_date else 'N/A',
                'Closure Remarks': getattr(item, 'closure_remarks', 'N/A')
            })

        return {
            'data': data,
            'headers': headers,
            'title': 'Vendor Records'
        }

    # [Keep all export methods unchanged from previous implementation]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Add any additional context variables here
        context.update(kwargs)
        return context


    # [Keep all your existing export methods exactly as they were]
    def export_to_csv(self, context):
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{context["title"].lower().replace(" ", "_")}.csv"'
        
        writer = csv.writer(response)
        writer.writerow(context['headers'])
        
        for row in context['data']:
            writer.writerow(row.values())
        
        return response

    def export_to_pdf(self, context):
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{context["title"].lower().replace(" ", "_")}.pdf"'
        
        buffer = BytesIO()
        p = canvas.Canvas(buffer)
        
        # PDF formatting
        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, 800, context["title"])
        p.setFont("Helvetica", 10)
        
        y = 750
        for row in context['data']:
            for header in context['headers']:
                p.drawString(50, y, f"{header}: {row.get(header, '')}")
                y -= 15
                if y < 50:
                    p.showPage()
                    y = 800
                    p.setFont("Helvetica-Bold", 14)
                    p.drawString(50, 800, context["title"])
                    p.setFont("Helvetica", 10)
            y -= 10  # Space between records
        
        p.save()
        pdf = buffer.getvalue()
        buffer.close()
        response.write(pdf)
        return response

    def export_to_excel(self, context):
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{context["title"].lower().replace(" ", "_")}.xlsx"'

        wb = Workbook()
        ws = wb.active
        ws.title = "Report"
        
        ws.append(context['headers'])
        
        for row in context['data']:
            ws.append(list(row.values()))
        
        wb.save(response)
        return response

    def export_to_word(self, context):
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{context["title"].lower().replace(" ", "_")}.docx"'

        document = Document()
        document.add_heading(context['title'], 0)
        
        for row in context['data']:
            for header in context['headers']:
                document.add_paragraph(f"{header}: {row.get(header, '')}")
            document.add_page_break()
        
        document.save(response)
        return response

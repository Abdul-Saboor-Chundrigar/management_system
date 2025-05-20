from django.shortcuts import render, HttpResponse
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from django.utils import timezone
from django.db.models import Subquery, OuterRef
from .models import UserLocation
import csv
from django.db import transaction

@login_required
# live/views.py
def admin_location_monitor(request):
    # Get active sessions (last 2 hours)
    active_sessions = UserLocation.objects.filter(
        timestamp__gte=timezone.now()-timezone.timedelta(hours=2)
    ).order_by('-timestamp').select_related('user')
    
    context = {
        'user_locations': active_sessions,
        'current_time': timezone.now()
    }
    return render(request, 'live/monitor.html', context)

@login_required
def refresh_locations(request):
    if request.method == 'POST':
        from .middleware import LocationTrackingMiddleware
        middleware = LocationTrackingMiddleware(get_response=lambda r: None)
        middleware(request)
        return JsonResponse({'status': 'success'})

@login_required
def delete_location(request, pk):
    if request.method == 'POST':
        UserLocation.objects.filter(id=pk).delete()
        return JsonResponse({'status': 'success'})

@login_required
def export_locations(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="user_locations.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['User', 'IP', 'Latitude', 'Longitude', 'City', 'Country', 'Timestamp'])
    
    for loc in UserLocation.objects.all():
        writer.writerow([
            loc.user.username,
            loc.ip,
            loc.latitude,
            loc.longitude,
            loc.city,
            loc.country,
            loc.timestamp
        ])
    
    return response

@login_required
def cleanup_locations(request):
    if request.method == 'POST':
        from datetime import timedelta
        cutoff = timezone.now() - timedelta(days=30)
        count = UserLocation.objects.filter(timestamp__lt=cutoff).delete()[0]
        return JsonResponse({'status': 'success', 'deleted': count})

@login_required
def update_location(request):
    if request.method == 'POST':
        try:
            latitude = request.POST.get('latitude')
            longitude = request.POST.get('longitude')
            
            if latitude and longitude:
                UserLocation.objects.create(
                    user=request.user,
                    ip=request.META.get('REMOTE_ADDR'),
                    latitude=latitude,
                    longitude=longitude,
                    city="Manual Update",
                    country="Unknown",
                    nearby_landmark="Browser Geolocation"
                )
                return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})
    
    return JsonResponse({'status': 'error', 'message': 'Invalid request'})

from django.http import HttpResponse
import csv
from reportlab.pdfgen import canvas
from docx import Document
from io import BytesIO

def export_locations(request, format=None):
    # Get format from URL or default to CSV
    format_type = format or request.GET.get('format', 'csv')
    
    if format_type == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="user_locations.csv"'
        
        writer = csv.writer(response)
        writer.writerow(['User', 'IP', 'City', 'Country', 'Latitude', 'Longitude', 'Timestamp'])
        
        for loc in UserLocation.objects.all():
            writer.writerow([
                loc.user.username,
                loc.ip,
                loc.city,
                loc.country,
                loc.latitude,
                loc.longitude,
                loc.timestamp.strftime('%Y-%m-%d %H:%M:%S')
            ])
        return response

    elif format_type == 'pdf':
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="user_locations.pdf"'
        
        p = canvas.Canvas(response)
        p.drawString(100, 800, "User Location Report")
        
        y = 750
        for loc in UserLocation.objects.all():
            p.drawString(100, y, f"{loc.user.username} - {loc.city}, {loc.country}")
            y -= 20
            if y < 50:
                p.showPage()
                y = 750
        p.save()
        return response

    elif format_type == 'doc':
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="user_locations.docx"'
        
        document = Document()
        document.add_heading('User Location Report', 0)
        
        table = document.add_table(rows=1, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'User'
        hdr_cells[1].text = 'IP'
        hdr_cells[2].text = 'City'
        hdr_cells[3].text = 'Country'
        hdr_cells[4].text = 'Latitude'
        hdr_cells[5].text = 'Longitude'
        hdr_cells[6].text = 'Timestamp'
        
        for loc in UserLocation.objects.all():
            row_cells = table.add_row().cells
            row_cells[0].text = loc.user.username
            row_cells[1].text = loc.ip
            row_cells[2].text = loc.city
            row_cells[3].text = loc.country
            row_cells[4].text = str(loc.latitude)
            row_cells[5].text = str(loc.longitude)
            row_cells[6].text = loc.timestamp.strftime('%Y-%m-%d %H:%M:%S')
        
        document.save(response)
        return response

    return HttpResponse("Invalid export format", status=400)
